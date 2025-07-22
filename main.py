import streamlit as st
import pandas as pd
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import base64

# Page configuration
st.set_page_config(
    page_title="Vibro-Sens Inspection System",
    page_icon="🏭",
    layout="wide"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        text-align: center;
        color: #2E86C1;
        margin-bottom: 30px;
    }
    .section-header {
        background-color: #E8F4FD;
        padding: 10px;
        border-radius: 5px;
        margin: 15px 0;
        color: #1B4F72;
        font-weight: bold;
    }
    .inspection-item {
        margin: 10px 0;
        padding: 5px;
        border-left: 3px solid #3498DB;
        padding-left: 15px;
    }
    .logo-container {
        display: flex;
        justify-content: center;
        margin-bottom: 30px;
    }
</style>
""", unsafe_allow_html=True)

def create_logo_placeholder():
    """Create a placeholder for the Ambatovy logo"""
    st.markdown("""
    <div class="logo-container">
        <div style="background-color: #E8F4FD; padding: 20px; border-radius: 10px; text-align: center; border: 2px solid #3498DB;">
            <h2 style="color: #1B4F72; margin: 0;">VIBROSens</h2>
            <p style="margin: 5px 0 0 0; color: #5D6D7E;">Condition Monitoring Rotating Equipment</p>
        </div>
    </div>
    """, unsafe_allow_html=True)

def initialize_session_state():
    """Initialize session state variables"""
    if 'inspection_data' not in st.session_state:
        st.session_state.inspection_data = {}
    if 'completed_sections' not in st.session_state:
        st.session_state.completed_sections = set()

def create_inspection_form():
    """Create the main inspection form"""
    
    # Header information
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Inspector Information")
        technician_name = st.text_input("Technician Name", placeholder="e.g., Rodin")
        group = st.text_input("Group", placeholder="e.g., Group A")
        
    with col2:
        st.subheader("Inspection Details")
        inspection_date = st.date_input("Inspection Date", datetime.now())
        equipment_tag = st.text_input("Equipment Tag #", value="31 - TM -")
        wo_number = st.text_input("Work Order #", placeholder="WO#")
    
    # Inspection type
    inspection_type = st.selectbox(
        "Select Inspection Type",
        ["Thickener I Rake Drive Hydraulic Power Pack", 
         "Thickener II Rake Drive Hydraulic Power Pack"]
    )
    
    # Visual and Vibration checkboxes
    col1, col2 = st.columns(2)
    with col1:
        visual_check = st.checkbox("Visual Inspection", value=True)
    with col2:
        vibration_check = st.checkbox("Vibration Check", value=True)
    
    return {
        'technician_name': technician_name,
        'group': group,
        'inspection_date': inspection_date,
        'equipment_tag': equipment_tag,
        'wo_number': wo_number,
        'inspection_type': inspection_type,
        'visual_check': visual_check,
        'vibration_check': vibration_check
    }

def safety_section():
    """Safety inspection section"""
    st.markdown('<div class="section-header">🔒 Safety</div>', unsafe_allow_html=True)
    
    with st.expander("Safety Inspection Items", expanded=True):
        safety_data = {}
        
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Equipment Tags:**")
            safety_data['equipment_tags'] = st.radio(
                "Equipment Tags Status", 
                ["OK", "Not OK"], 
                key="safety_tags"
            )
            
            st.markdown("**Hand Rail/Grating:**")
            safety_data['handrail_grating'] = st.radio(
                "Hand Rail/Grating Status", 
                ["OK", "Not OK"], 
                key="safety_handrail"
            )
            
        with col2:
            st.markdown("**Housekeeping - Cleaning:**")
            safety_data['housekeeping'] = st.radio(
                "Housekeeping Status", 
                ["OK", "Not OK"], 
                key="safety_housekeeping"
            )
            
            st.markdown("**Terminal Box/Grounding Cables:**")
            safety_data['terminal_grounding'] = st.radio(
                "Terminal Box/Grounding Status", 
                ["OK", "Not OK"], 
                key="safety_terminal"
            )
        
        safety_data['comments'] = st.text_area("Safety Comments", key="safety_comments")
        
    return safety_data

def general_rake_condition_section():
    """General rake operating condition section"""
    st.markdown('<div class="section-header">⚙️ General Rake Operating Condition</div>', unsafe_allow_html=True)
    
    with st.expander("Operating Parameters", expanded=True):
        operating_data = {}
        
        col1, col2 = st.columns(2)
        with col1:
            operating_data['drive_oil_pressure'] = st.number_input(
                "Drive Hydraulic Supply Oil Pressure (MPa)", 
                min_value=0.0, max_value=50.0, step=0.1,
                key="drive_pressure"
            )
            
            operating_data['rake_torque_pressure'] = st.number_input(
                "Rake Torque Pressure (MPa)", 
                min_value=0.0, max_value=50.0, step=0.1,
                key="torque_pressure"
            )
            
        with col2:
            operating_data['rake_lift_pressure'] = st.number_input(
                "Rake Lift Pressure (MPa) - Target: 9-10 MPa", 
                min_value=0.0, max_value=15.0, step=0.1,
                key="lift_pressure"
            )
            
            if operating_data['rake_lift_pressure'] < 9 or operating_data['rake_lift_pressure'] > 10:
                st.warning("⚠️ Rake lift pressure is outside normal range (9-10 MPa)")
        
    return operating_data

def reservoir_section():
    """Reservoir inspection section"""
    st.markdown('<div class="section-header">🛢️ Reservoir</div>', unsafe_allow_html=True)
    
    with st.expander("Reservoir Inspection Items", expanded=True):
        reservoir_data = {}
        
        # Temperature readings
        col1, col2, col3 = st.columns(3)
        with col1:
            reservoir_data['prv1_temp'] = st.number_input("PRV 1 Temperature (°C)", key="prv1_temp")
        with col2:
            reservoir_data['prv2_temp'] = st.number_input("PRV 2 Temperature (°C)", key="prv2_temp")
        with col3:
            reservoir_data['prv3_temp'] = st.number_input("PRV 3 Temperature (°C)", key="prv3_temp")
        
        # Delta pressure
        reservoir_data['delta_pressure'] = st.number_input(
            "Delta Pressure Across Filter (kPa) - Target: < 300 kPa", 
            min_value=0.0, step=1.0,
            key="delta_pressure"
        )
        
        if reservoir_data['delta_pressure'] >= 300:
            st.warning("⚠️ Delta pressure is above recommended limit (300 kPa)")
        
        # Inspection items
        inspection_items = [
            ("oil_leaks", "Check hydraulic oil reservoir for oil leaks"),
            ("condensate", "Check hydraulic oil reservoir for condensate built up"),
            ("contamination", "Check hydraulic oil for contamination (dirty/milky)"),
            ("panel_fittings", "Check instrument and fittings on panel for oil leaks"),
            ("breather_condition", "Check reservoir breather condition")
        ]
        
        for key, description in inspection_items:
            st.markdown(f"**{description}:**")
            reservoir_data[key] = st.radio(
                description, 
                ["OK", "Not OK"], 
                key=f"reservoir_{key}"
            )
        
        # Filter color indicator
        st.markdown("**Filter Color Indicator:**")
        reservoir_data['filter_color'] = st.radio(
            "Filter Color Status",
            ["Green (OK)", "Yellow (Dirty)", "Red (Bypass)"],
            key="filter_color"
        )
        
        reservoir_data['comments'] = st.text_area("Reservoir Comments", key="reservoir_comments")
        
    return reservoir_data

def hydraulic_drive_unit_section():
    """Hydraulic drive unit section"""
    st.markdown('<div class="section-header">🔧 Hydraulic Drive Unit</div>', unsafe_allow_html=True)
    
    with st.expander("Hydraulic Drive Unit Inspection", expanded=True):
        drive_data = {}
        
        # Temperature readings
        col1, col2 = st.columns(2)
        with col1:
            drive_data['nde_temp'] = st.number_input("NDE Temperature (°C)", key="drive_nde_temp")
        with col2:
            drive_data['motor_body_temp'] = st.number_input("Motor Body Temperature (°C)", key="drive_motor_temp")
        
        # Inspection items
        inspection_items = [
            ("general_condition", "General condition & Noise"),
            ("hold_down_bolts", "Hold down bolts and Foundation base plate"),
            ("cooling_lube", "Cooling system and Lube fitting integrity")
        ]
        
        for key, description in inspection_items:
            st.markdown(f"**{description}:**")
            drive_data[key] = st.radio(
                description, 
                ["OK", "Not OK"], 
                key=f"drive_{key}"
            )
        
        drive_data['comments'] = st.text_area("Hydraulic Drive Unit Comments", key="drive_comments")
        
    return drive_data

def hydraulic_pump_section():
    """Hydraulic oil supply pump section"""
    st.markdown('<div class="section-header">🔄 Hydraulic Oil Supply Pump</div>', unsafe_allow_html=True)
    
    with st.expander("Pump Inspection Items", expanded=True):
        pump_data = {}
        
        # Temperature reading
        pump_data['pump_temp'] = st.number_input("Pump Temperature (°C)", key="pump_temp")
        
        # Inspection items
        inspection_items = [
            ("general_condition", "General condition & Noise"),
            ("pedestal_bolts", "Pedestal hold down bolts and Foundation base plate"),
            ("casing_fittings", "Pump casing & Suction/discharge line fittings"),
            ("flexible_hoses", "Check Flexible hose supply lines for chafe and cracks")
        ]
        
        for key, description in inspection_items:
            st.markdown(f"**{description}:**")
            pump_data[key] = st.radio(
                description, 
                ["OK", "Not OK"], 
                key=f"pump_{key}"
            )
        
        pump_data['comments'] = st.text_area("Pump Comments", key="pump_comments")
        
    return pump_data
'''
def create_docx_report(inspection_info, inspection_data):
    """Create a Word document report"""
    doc = Document()
    
    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "AMBATOVY - Condition Monitoring Rotating Equipment"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Title
    title = doc.add_heading('Thickener Hydraulic Power Pack CM Check Sheet', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Inspection Information
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Check by:", f"{inspection_info['technician_name']} / {inspection_info['group']}"),
        ("Date:", inspection_info['inspection_date'].strftime("%d/%m/%Y")),
        ("Equipment Tag #:", inspection_info['equipment_tag']),
        ("Work Order #:", inspection_info['wo_number']),
        ("Visual Check:", "✓" if inspection_info['visual_check'] else "✗"),
        ("Vibration Check:", "✓" if inspection_info['vibration_check'] else "✗")
    ]
    
    for i, (key, value) in enumerate(info_data):
        info_table.cell(i, 0).text = key
        info_table.cell(i, 1).text = str(value)
    
    # Safety Section
    doc.add_heading('Safety', level=1)
    safety_table = doc.add_table(rows=5, cols=2)
    safety_table.style = 'Table Grid'
    
    safety_items = [
        ("Equipment Tags:", inspection_data.get('safety', {}).get('equipment_tags', '')),
        ("Hand Rail/Grating:", inspection_data.get('safety', {}).get('handrail_grating', '')),
        ("Housekeeping:", inspection_data.get('safety', {}).get('housekeeping', '')),
        ("Terminal Box/Grounding:", inspection_data.get('safety', {}).get('terminal_grounding', '')),
        ("Comments:", inspection_data.get('safety', {}).get('comments', ''))
    ]
    
    for i, (item, status) in enumerate(safety_items):
        safety_table.cell(i, 0).text = item
        safety_table.cell(i, 1).text = str(status)
    
    # Operating Conditions
    doc.add_heading('General Rake Operating Condition', level=1)
    operating_table = doc.add_table(rows=3, cols=2)
    operating_table.style = 'Table Grid'
    
    operating_items = [
        ("Drive Oil Pressure (MPa):", inspection_data.get('operating', {}).get('drive_oil_pressure', '')),
        ("Rake Torque Pressure (MPa):", inspection_data.get('operating', {}).get('rake_torque_pressure', '')),
        ("Rake Lift Pressure (MPa):", inspection_data.get('operating', {}).get('rake_lift_pressure', ''))
    ]
    
    for i, (item, value) in enumerate(operating_items):
        operating_table.cell(i, 0).text = item
        operating_table.cell(i, 1).text = str(value)
    
    # Add other sections as needed...
    
    return doc
'''
#-------------------------------------------Nouveau----------------------------------------------------------------------------------------------------

def create_docx_report(inspection_info, inspection_data):
    """Create a comprehensive Word document report aligned with provided templates"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # Header with logo placeholder
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_run = header_para.add_run()
    # In a real implementation, you would add an actual image here
    # header_run.add_picture("logo.png", width=Inches(1.0))
    header_run.add_text("AMBATOVY - Condition Monitoring Rotating Equipment")
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title
    title = doc.add_heading('Thickener Hydraulic Power Pack CM Check Sheet', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(14)
    title.style.font.bold = True

    # Inspection Information
    doc.add_heading('Inspection Details', level=1)
    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = 'Table Grid'
    info_table.autofit = False
    
    # Set column widths
    for cell in info_table.columns[0].cells:
        cell.width = Inches(1.2)
    for cell in info_table.columns[1].cells:
        cell.width = Inches(2.0)
    for cell in info_table.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in info_table.columns[3].cells:
        cell.width = Inches(2.0)
    
    # Populate header row
    hdr_cells = info_table.rows[0].cells
    hdr_cells[0].text = "Check by:"
    hdr_cells[1].text = f"{inspection_info['technician_name']} / {inspection_info['group']}"
    hdr_cells[2].text = "Date:"
    hdr_cells[3].text = inspection_info['inspection_date'].strftime("%d/%m/%Y")
    
    # Populate data rows
    rows = [
        ("Review by:", ""),
        ("Equipment Tag #:", inspection_info['equipment_tag']),
        ("Work Order #:", inspection_info['wo_number']),
        ("Visual Check:", "✓" if inspection_info['visual_check'] else "✗"),
        ("Vibration Check:", "✓" if inspection_info['vibration_check'] else "✗")
    ]
    
    for i, (label, value) in enumerate(rows, start=1):
        row_cells = info_table.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = value
        if i < len(rows):
            row_cells[2].text = ""
            row_cells[3].text = ""

    # Safety Section
    safety = inspection_data.get('safety', {})
    doc.add_heading('Safety', level=1)
    safety_table = doc.add_table(rows=6, cols=2)
    safety_table.style = 'Table Grid'
    
    safety_items = [
        ("Equipment Tags:", safety.get('equipment_tags', '')),
        ("Hand Rail/Grating:", safety.get('handrail_grating', '')),
        ("Housekeeping:", safety.get('housekeeping', '')),
        ("Terminal Box/Grounding:", safety.get('terminal_grounding', '')),
        ("Coupling Guard:", safety.get('coupling_guard', '')),
        ("Comments:", safety.get('comments', ''))
    ]
    
    for i, (item, status) in enumerate(safety_items):
        row_cells = safety_table.rows[i].cells
        row_cells[0].text = item
        row_cells[1].text = str(status)
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    # General Rake Operating Condition
    operating = inspection_data.get('operating', {})
    doc.add_heading('General Rake Operating Condition', level=1)
    doc.add_paragraph("Check rake drive system and record the following data.")
    
    operating_table = doc.add_table(rows=8, cols=2)
    operating_table.style = 'Table Grid'
    
    operating_items = [
        ("Drive Oil Pressure (MPa):", operating.get('drive_oil_pressure', '')),
        ("Rake Torque Pressure (MPa):", operating.get('rake_torque_pressure', '')),
        ("Rake Lift Pressure (MPa):", operating.get('rake_lift_pressure', '')),
        ("Rake Lift Pressure While Lifting (MPa):", operating.get('rake_lift_pressure_lifting', '')),
        ("Rake Lift Pressure While Lowering (MPa):", operating.get('rake_lift_pressure_lowering', '')),
        ("Thickener Rake Position:", operating.get('rake_position', '')),
        ("Thickener Rake Torque:", operating.get('rake_torque', '')),
        ("Thickener Rake Speed:", operating.get('rake_speed', ''))
    ]
    
    for i, (item, value) in enumerate(operating_items):
        row_cells = operating_table.rows[i].cells
        row_cells[0].text = item
        row_cells[1].text = str(value)
        row_cells[0].paragraphs[0].runs[0].font.bold = True

    # Reservoir Section
    reservoir = inspection_data.get('reservoir', {})
    doc.add_heading('Reservoir', level=1)
    doc.add_paragraph("Check hydraulic oil reservoir and record the following data.")
    
    reservoir_table = doc.add_table(rows=11, cols=4)
    reservoir_table.style = 'Table Grid'
    
    # Reservoir items
    reservoir_items = [
        ("Check hydraulic oil reservoir for oil leaks", reservoir.get('oil_leaks', ''), "", ""),
        ("Check hydraulic oil reservoir for condensate built up", reservoir.get('condensate', ''), "", ""),
        ("Check hydraulic oil for contamination (dirty/milky)", reservoir.get('contamination', ''), "", ""),
        ("Check instrument and fittings on panel for oil leaks", reservoir.get('panel_fittings', ''), "", ""),
        ("Check reservoir breather condition", reservoir.get('breather_condition', ''), "", ""),
        ("Delta Pressure Across Filter (kPa)", "", reservoir.get('delta_pressure', ''), ""),
        ("Filter Color Indicator", reservoir.get('filter_color', ''), "", ""),
        ("PRV 1 Temperature (°C)", "", "", reservoir.get('prv1_temp', '')),
        ("PRV 2 Temperature (°C)", "", "", reservoir.get('prv2_temp', '')),
        ("PRV 3 Temperature (°C)", "", "", reservoir.get('prv3_temp', '')),
        ("Comments:", reservoir.get('comments', ''), "", "")
    ]
    
    for i, row_data in enumerate(reservoir_items):
        row_cells = reservoir_table.rows[i].cells
        for j, cell_value in enumerate(row_data):
            row_cells[j].text = str(cell_value)
            if j == 0 and i < 5:  # Make inspection items bold
                row_cells[j].paragraphs[0].runs[0].font.bold = True

    # Hydraulic Drive Unit Section
    drive = inspection_data.get('hydraulic_drive', {})
    doc.add_heading('Hydraulic Drive Unit', level=1)
    
    drive_table = doc.add_table(rows=6, cols=5)
    drive_table.style = 'Table Grid'
    
    # Header row
    hdr_cells = drive_table.rows[0].cells
    hdr_cells[0].text = "Inspection Item"
    hdr_cells[1].text = "Status"
    hdr_cells[3].text = "Measurements"
    
    # Drive unit items
    drive_items = [
        ("General condition & Noise", drive.get('general_condition', ''), "NDE Temperature (°C)", drive.get('nde_temp', '')),
        ("Hold down bolts and Foundation base plate", drive.get('hold_down_bolts', ''), "NDE Temperature (°C)", drive.get('nde_temp2', '')),
        ("Cooling system and Lube fitting integrity", drive.get('cooling_lube', ''), "Motor Body Temperature (°C)", drive.get('motor_body_temp', '')),
        ("", "", "Vibration (mm/sec)", drive.get('vibration', '')),
        ("", "", "Temperature (°C)", drive.get('temperature', '')),
        ("Comments:", drive.get('comments', ''), "", "")
    ]
    
    for i, row_data in enumerate(drive_items, start=1):
        row_cells = drive_table.rows[i].cells
        for j, cell_value in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = str(cell_value)
                if i == 5:  # Comments row
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].merge(row_cells[2])
                    row_cells[0].merge(row_cells[3])
                    row_cells[0].merge(row_cells[4])

    # Hydraulic Pump Section
    pump = inspection_data.get('hydraulic_pump', {})
    doc.add_heading('Hydraulic Oil Supply Pump', level=1)
    
    pump_table = doc.add_table(rows=9, cols=5)
    pump_table.style = 'Table Grid'
    
    # Header row
    hdr_cells = pump_table.rows[0].cells
    hdr_cells[0].text = "Inspection Item"
    hdr_cells[1].text = "Status"
    hdr_cells[3].text = "Measurements"
    
    # Pump items
    pump_items = [
        ("General condition & Noise", pump.get('general_condition', ''), "Pump Temperature (°C)", pump.get('pump_temp', '')),
        ("Pedestal hold down bolts and Foundation base plate", pump.get('pedestal_bolts', ''), "Vibration (mm/sec)", pump.get('vibration', '')),
        ("Pump casing & Suction/discharge line fittings", pump.get('casing_fittings', ''), "Temperature (°C)", pump.get('temperature', '')),
        ("Check Flexible hose supply lines for chafe and cracks", pump.get('flexible_hoses', ''), "", ""),
        ("", "", "", ""),
        ("Comments:", pump.get('comments', ''), "", "")
    ]
    
    for i, row_data in enumerate(pump_items, start=1):
        row_cells = pump_table.rows[i].cells
        for j, cell_value in enumerate(row_data):
            if j < len(row_cells):
                row_cells[j].text = str(cell_value)
                if i == 5:  # Comments row
                    row_cells[0].merge(row_cells[1])
                    row_cells[0].merge(row_cells[2])
                    row_cells[0].merge(row_cells[3])
                    row_cells[0].merge(row_cells[4])

    # Add more sections similarly for:
    # - Motor M2
    # - Rotary Hydraulic Drive Motor and Gearbox
    # - Planetary Gear Reducer
    # - Rake Lift mechanism
    # - Automatic Grease Lubrication Unit
    
    return doc

# ... (le reste du code reste inchangé)

#----------------------------------------------Fin-----------------------------------------------------------------------------------------------------
def export_to_csv(inspection_info, inspection_data):
    """Export inspection data to CSV"""
    csv_data = []
    
    # Basic information
    base_info = {
        'Date': inspection_info['inspection_date'].strftime("%Y-%m-%d"),
        'Technician': inspection_info['technician_name'],
        'Group': inspection_info['group'],
        'Equipment_Tag': inspection_info['equipment_tag'],
        'WO_Number': inspection_info['wo_number'],
        'Inspection_Type': inspection_info['inspection_type'],
        'Visual_Check': inspection_info['visual_check'],
        'Vibration_Check': inspection_info['vibration_check']
    }
    
    # Flatten inspection data
    flattened_data = {}
    for section, data in inspection_data.items():
        if isinstance(data, dict):
            for key, value in data.items():
                flattened_data[f"{section}_{key}"] = value
        else:
            flattened_data[section] = data
    
    # Combine all data
    combined_data = {**base_info, **flattened_data}
    csv_data.append(combined_data)
    
    df = pd.DataFrame(csv_data)
    return df

def main():
    """Main application function"""
    initialize_session_state()
    
    # Display logo
    create_logo_placeholder()
    
    st.markdown('<h1 class="main-header">Hydraulic Power Pack Inspection System</h1>', 
                unsafe_allow_html=True)
    
    # Create main form
    with st.form("inspection_form"):
        # Get inspection info
        inspection_info = create_inspection_form()
        
        st.markdown("---")
        
        # Inspection sections
        inspection_data = {}
        
        # Safety section
        inspection_data['safety'] = safety_section()
        
        # General rake operating condition
        inspection_data['operating'] = general_rake_condition_section()
        
        # Reservoir section
        inspection_data['reservoir'] = reservoir_section()
        
        # Hydraulic drive unit
        inspection_data['hydraulic_drive'] = hydraulic_drive_unit_section()
        
        # Hydraulic pump
        inspection_data['hydraulic_pump'] = hydraulic_pump_section()
        
        # Form submission
        st.markdown("---")
        submitted = st.form_submit_button("Complete Inspection", type="primary")
        
        if submitted:
            if not inspection_info['technician_name'] or not inspection_info['group']:
                st.error("Please enter technician name and group before submitting.")
            else:
                st.success("Inspection completed successfully!")
                
                # Store data in session state
                st.session_state.inspection_info = inspection_info
                st.session_state.inspection_data = inspection_data
    
    # Export options (only show if inspection is completed)
    if hasattr(st.session_state, 'inspection_info'):
        st.markdown("---")
        st.subheader("📥 Export Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 Download Word Report"):
                doc = create_docx_report(st.session_state.inspection_info, 
                                       st.session_state.inspection_data)
                
                # Save document to bytes
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                st.download_button(
                    label="Download DOCX",
                    data=doc_bytes.getvalue(),
                    file_name=f"inspection_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col2:
            if st.button("📊 Download CSV Data"):
                df = export_to_csv(st.session_state.inspection_info, 
                                 st.session_state.inspection_data)
                
                csv_buffer = io.StringIO()
                df.to_csv(csv_buffer, index=False)
                csv_data = csv_buffer.getvalue()
                
                st.download_button(
                    label="Download CSV",
                    data=csv_data,
                    file_name=f"inspection_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )

if __name__ == "__main__":
    main()
